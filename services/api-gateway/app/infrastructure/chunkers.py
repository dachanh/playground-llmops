from __future__ import annotations

import hashlib
from pathlib import Path
from typing import List

import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader


def load_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if path.suffix.lower() == ".docx":
        doc = docx.Document(str(path))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return path.read_text(encoding="utf-8", errors="ignore")


def chunk_text(text: str, source: str, chunk_size: int, overlap: int) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    chunks = splitter.split_text(text)
    documents: List[Document] = []
    for idx, chunk in enumerate(chunks):
        chunk_id = f"{source}-{idx}"
        chunk_hash = hashlib.sha256(chunk.encode("utf-8")).hexdigest()
        metadata = {"source": source, "chunk_id": chunk_id, "hash": chunk_hash}
        documents.append(Document(page_content=chunk, metadata=metadata))
    return documents
